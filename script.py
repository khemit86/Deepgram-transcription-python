from deepgram import Deepgram
import asyncio, json
import datetime
import docx
from docx.shared import Pt
from docx import Document



DEEPGRAM_API_KEY = 'xxxxxxxxxxxxxxxxxxxxxxxx'
PATH_TO_FILE = 'SERVICE_1_TEST1.mov'



async def main():
    # Initializes the Deepgram SDK
    deepgram = Deepgram(DEEPGRAM_API_KEY)
    # Open the audio file
    with open(PATH_TO_FILE, 'rb') as audio:
        # ...or replace mimetype as appropriate
        source = {'buffer': audio, 'mimetype': 'video/mov'}
        #response = await deepgram.transcription.prerecorded(source, {'punctuate': True,'diarize':True,'paragraphs':True,'tier':'enhanced'})
        response = await deepgram.transcription.prerecorded(source, {'punctuate': True,'diarize':True,'utterances':True})
        
        #print(response['results']['channels'][0]['alternatives'][0]['paragraphs'])
        #print(response['results']['utterances'])
        #print(json.dumps(response['utterances']))
        document = Document()
        #transcript = ''
        style = document.styles['Normal']
        document.add_heading('Transcription Document', 0) 
        doc_para = document.add_paragraph('\n')

        font = style.font
        font.name = 'MS Gothic'
        font.size = Pt(10)    
       
        for data in response['results']['utterances']:
            #print(data['transcript'],data['speaker'])
            #transcript = 'Speaker'+str(data['speaker'])+': '+data['transcript']+'\n\n'
            #print(transcript)
            
            doc_para.add_run("["+str(datetime.timedelta(seconds=int(data['start'])))+"]\n")
            doc_para.add_run('Speaker'+str(data['speaker'])+': ').bold = True
            doc_para.add_run(data['transcript']+'\n\n')
        # Create an instance of a word document
        #p = doc.add_paragraph('A plain paragraph having some ')
        #p.add_run('bold').bold = True
        #p.add_run(' and some ')
        #p.add_run('italic.').italic = True
          
        # Add a Title to the document 
        #doc.add_heading('GeeksForGeeks', 0)
          
        # Adding paragraph with Increased font size
        #doc.add_heading('Increased Font Size Paragraph:', 2)

        #para = doc.add_paragraph().add_run(
        #    'GeeksforGeeks is a Computer Science portal for geeks.')
        # Increasing size of the font
        #para.font.size = Pt(12)
          
        # Adding paragraph with normal font size
        #doc.add_heading('Normal Font Size Paragraph:', 3)
        #doc.add_paragraph(
        #    'GeeksforGeeks is a Computer Science portal for geeks.')
        
        document.save('demo.docx')

asyncio.run(main())

